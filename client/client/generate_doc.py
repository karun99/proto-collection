from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Cover Page ──
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ROYAL MUTTON HUB')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Vijayawada-based B2B Mutton Distributor\nProject Documentation')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
datep = doc.add_paragraph()
datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = datep.add_run(f'Document Generated: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── Table of Contents ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Business Information',
    '3. Technology Stack',
    '4. Project Structure',
    '5. Features',
    '6. File Reference',
    '7. Configuration Guide',
    '8. Deployment Guide',
    '9. Admin Dashboard Guide',
    '10. User Manual — Admin Operations',
    '11. localStorage Reference',
    '12. Build Scripts',
    '13. Screenshot Checklist',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. Project Overview ──
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'Royal Mutton Hub is a fully static, no-backend website for a Vijayawada-based B2B mutton distributor. '
    'It serves hotels, restaurants, caterers, mess, and retail chains across Krishna, NTR, and Guntur districts '
    'with same-day delivery. The brand emphasizes fresh, pure, tender mutton delivered daily with a "No Refrigeration — Only Freshness" USP.'
)
doc.add_paragraph(
    'The project includes an AI-powered chatbot (via OpenRouter), a rich recipe database (RecipeDB) with chef-credited '
    'recipes and user submissions, lead capture, SEO-optimized static pages, a full admin dashboard, anti-bot protections, '
    'meeting scheduling, and password management — all running client-side with localStorage-based persistence.'
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Developed by: ')
run.bold = True
p.add_run('Guiding Key (http://guidingkey.com)')
p = doc.add_paragraph()
run = p.add_run('Client: ')
run.bold = True
p.add_run('Royal Mutton Hub, Vijayawada')

doc.add_heading('Key Differentiators', level=2)
bullets = [
    'No refrigeration — fresh mutton delivered daily',
    'Indigenous Vijayawada brand since 2005, rooted in Andhra Pradesh culinary heritage',
    'B2B-focused: dedicated service for hotels, restaurants, caterers, mess, retail chains',
    'Same-day delivery across Krishna, NTR & Guntur districts',
    'All prices in INR for budget-friendly positioning',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ── 2. Business Information ──
doc.add_heading('2. Business Information', level=1)

table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'
data = [
    ('Business Name', 'Royal Mutton Hub'),
    ('Type', 'B2B Mutton Distributor'),
    ('Address', '# 9-60-75, Ganapathi Rao Road, Kothapeta, Vijayawada - 520001'),
    ('Address (Telugu)', '♆ 9-60-75, గణపతి రావు రోడ్, కొత్తపేట, విజయవాడ - 520001'),
    ('Hours', 'Daily 6 AM – 8 PM'),
    ('Service Area', 'Krishna, NTR & Guntur Districts'),
    ('Email', 'rmhubrj18@gmail.com'),
    ('Delivery', 'Same-day delivery'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('B2B Services', level=2)
services = [
    'Hotel Supply — Regular bulk orders for hotel kitchens',
    'Restaurant Supply — Premium cuts for restaurant menus',
    'Caterer Supply — Event-based and volume orders for caterers',
    'Mess & Canteen — Daily supply for mess and canteens',
    'Retail Chain — Quality assured supply for retail stores',
    'Bulk Supply — Wholesale pricing for large volume buyers',
]
for s in services:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ── 3. Technology Stack ──
doc.add_heading('3. Technology Stack', level=1)

techs = [
    ('Frontend', 'Vanilla HTML5, CSS3, JavaScript (ES6+) — no frameworks or build tools'),
    ('Icons', 'Font Awesome 6.5.1 (CDN)'),
    ('Typography', 'Google Inter Font (CDN)'),
    ('AI Chatbot', 'OpenRouter API (configurable model, default: openrouter/free)'),
    ('Encryption', 'AES-256-GCM via Web Crypto API (client-side)'),
    ('Maps', 'OpenStreetMap embedded'),
    ('Storage', 'localStorage (client-side persistence)'),
    ('Deployment', 'Netlify (recommended) or Vercel'),
    ('Recipe SEO', 'JSON-LD structured data, Open Graph tags, canonical URLs'),
]

table = doc.add_table(rows=len(techs), cols=2)
table.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate(techs):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_page_break()

# ── 4. Project Structure ──
doc.add_heading('4. Project Structure', level=1)
code = """C:.
├── index.html              Main site (8 sections: Hero, How We Work, B2B Supply,
│                           Why Us, Testimonials, Recipes, Contact, Footer)
├── recipedb.html           Standalone recipe database page
├── admin.html              Admin dashboard (9 tabs)
├── utils.js                Shared constants (BUSINESS object), Crypto module,
│                           sanitization, helpers. Sets window.__RMH.API_KEY
├── embed.js                AI chatbot widget with tabs: Chat, Quote, Order,
│                           Meeting, Policy. OpenRouter integration.
├── chef-recipes.js         9 chef-credited recipes from 8 Indian chefs
├── generate-recipes.js     Build script — generates RecipeDB/*.html static pages
├── fetch-commodity-prices.js  Build script — fetches mandi prices from data.gov.in
├── logo.png                Brand logo
├── netlify.toml            Netlify deployment config (CSP, caching, 404)
├── vercel.json             Vercel deployment config (CSP, caching, 404)
├── package.json            Project metadata
├── generate_doc.py         This documentation generator
├── RecipeDB/               Generated static recipe pages (22 files)
│   ├── index.html          Recipe index with JSON-LD ItemList
│   ├── classic-mutton-biryani.html
│   ├── andhra-mutton-curry.html
│   ├── ... (20 more recipe pages)"""
for line in code.split('\n'):
    p = doc.add_paragraph(line)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)

doc.add_page_break()

# ── 5. Features ──
doc.add_heading('5. Features', level=1)

doc.add_heading('5.1 Main Website (index.html)', level=2)
features_main = [
    'Hero — Indigenous brand badges ("Indigenous Vijayawada Brand", "FSSAI Licensed", "B2B Since 2005"), CTA buttons (Enquire / WhatsApp / Call), district badges for Krishna, NTR & Guntur',
    'How We Work — 5-step B2B process: Enquire, Custom Quote, Order Processing, Fresh Preparation, Same-Day Delivery',
    'B2B Supply — 6 service cards targeting different client segments',
    'Why Us — Indigenous highlight box + 9 benefit cards (No Refrigeration, Same-Day Delivery, FSSAI Certified, etc.)',
    'Testimonials — 6 client reviews with name, rating, and feedback',
    'Recipe Database — Grid with search, filters (diet, budget, cuisine, meal type, category), submission form, modal with ratings, "I Cooked This!", text reviews, sharing',
    'Contact — Bilingual address (English/Telugu), OpenStreetMap embed, enquiry form with name/phone/email/message',
    'Footer — Indigenous brand tagline, Guiding Key credit, admin link',
]
for f in features_main:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.2 AI Chatbot (embed.js)', level=2)
features_bot = [
    'OpenRouter API integration with configurable model (reads rmh_admin_model from localStorage)',
    'Canned response fallback when API is unavailable',
    'Two persona modes: Sales (default) and Executive',
    'Prompt injection detection',
    'Lead capture to rmh_analytics localStorage',
    'Meeting scheduling tab with date/time picker',
    'Quote request, Order placement, Policy info tabs',
    'Rich message formatting: **bold**, *italic*, bullet lists, paragraphs',
    'CTA buttons for quick replies',
    'Toast notifications instead of alerts',
]
for f in features_bot:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.3 Recipe Database (RecipeDB/)', level=2)
features_recipe = [
    '22 static HTML pages (12 seed recipes, 9 chef recipes, 1 index)',
    'Per-recipe JSON-LD structured data (Recipe schema)',
    'Open Graph tags, Twitter cards, canonical URLs for SEO',
    'Chef credit badges with source links',
    'Social sharing buttons (WhatsApp, Facebook, X, Telegram, Email, Copy Link)',
    'Indigenous brand badge on every page',
    'Prev/Next navigation between recipes',
    'Guiding Key footer link',
]
for f in features_recipe:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.4 Recipe Features', level=2)
features_rec = [
    'Questionnaire form with input sanitization (strip HTML + encode special chars)',
    'Rate limiting: max 3 submissions per hour per browser',
    'Filters: diet (veg/non-veg), budget (low/medium/high), cuisine (Andhra/Hyderabadi/North Indian/Mughlai), meal type, category',
    'Star rating (1-5) stored per recipe ID',
    '"I Cooked This!" counter with appreciation tracking',
    'Text reviews with name, date, and content',
    'Average rating display on modal',
    'Improved circular icon-only share buttons with hover effects',
    'Emoji-based image placeholders (no hosting needed)',
    'Cost estimates in INR',
    'Equipment badges for each recipe',
]
for f in features_rec:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.5 Admin Dashboard (admin.html)', level=2)
doc.add_paragraph('Password-protected dashboard (default: admin123) with AES-256-GCM encryption. 9 tabs:')

tabs = [
    ('Overview', 'Stats cards (total recipes, API calls, leads, meetings), recent activity feed'),
    ('API Keys', 'Encrypted OpenRouter key storage, show/hide toggle, test validation, model selector, rotate functionality. Saves plaintext copy to rmh_bot_apikey for chatbot.'),
    ('Password', 'Change admin password with current/new/confirm validation, displays current hash'),
    ('Branding', 'Edit business name, tagline, address (EN/TE), phone, WhatsApp, email. Mutton cuts CRUD (add/edit/delete). Custom CSS injection.'),
    ('Theming', 'Dark/light mode toggle, accent color picker, font selector, border radius slider, 5 preset themes'),
    ('RecipeDB', 'View/search all recipes, add/edit/delete user & chef recipes, bulk delete user recipes, export/import JSON'),
    ('Meetings', 'View/search/filter B2B meeting requests by status. Actions: Confirm, Complete, Cancel, Reopen. Status flow: pending -> confirmed -> completed, or cancelled.'),
    ('Data Management', 'Encrypted backup export/import with password. Local snapshot backups (max 20). Danger zone: clear user recipes, reset analytics, factory reset. CSV/JSON export/import for recipes, meetings, leads, reviews.'),
    ('API Usage', 'Stats: page views, API calls, leads. Leads table with source/date. API calls history with model and query preview.'),
]
for name, desc in tabs:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('5.6 Security Features', level=2)
sec = [
    'AES-256-GCM encryption for API keys, backups, and sensitive data',
    'DevTools detector: checks outerWidth - innerWidth > 200 || outerHeight - innerHeight > 200 every 2s; blanks page if detected',
    'Right-click disabled, text selection disabled, copy/cut blocked, drag blocked',
    'Keyboard shortcut blocking: Ctrl+C/U/S/P/I/J/Y/A/F/E/H, F12',
    'user-select: none CSS on all 3 HTML pages',
    'Email obfuscation: rmhubrj18@gmail.com split into JS variables, decoded at runtime',
    'Input sanitization on all user-submitted data',
    'Rate limiting on recipe submissions',
    'CSP headers via meta tag, Netlify TOML, and Vercel JSON',
]
for s in sec:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ── 6. File Reference ──
doc.add_heading('6. File Reference', level=1)

doc.add_heading('6.1 index.html', level=2)
doc.add_paragraph('Main entry page with all 8 sections. Script loading order:')
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
for i, (s, a, p) in enumerate([
    ('chef-recipes.js', 'defer', 'Chef recipe data (28KB)'),
    ('utils.js', 'sync', 'BUSINESS constants, crypto, helpers'),
    ('embed.js', 'defer', 'Chatbot widget'),
    ('inline script', 'sync', 'DOMContentLoaded initialization'),
]):
    table.rows[i].cells[0].text = s
    table.rows[i].cells[1].text = a
    table.rows[i].cells[2].text = p

doc.add_paragraph('All dynamic content renders inside DOMContentLoaded after deferred scripts execute.')

doc.add_heading('6.2 utils.js', level=2)
fields = [
    ('BUSINESS.recipes[]', '12 seed recipes with full metadata'),
    ('BUSINESS.services[]', '6 B2B service types'),
    ('BUSINESS.cuts[]', '10 mutton product strings'),
    ('BUSINESS.districts[]', '3 service districts'),
    ('BUSINESS.addressTelugu', 'Telugu address string'),
    ('BUSINESS.indigenousNote', 'Indigenous brand messaging'),
    ('mergeChefRecipes()', 'Auto-merges chef recipes into main recipe array'),
    ('Crypto.encrypt()', 'AES-256-GCM encrypt with password-derived key'),
    ('Crypto.decrypt()', 'AES-256-GCM decrypt with password-derived key'),
    ('safeParse()', 'Safe JSON.parse with fallback'),
    ('escapeHtml()', 'HTML entity encoding'),
    ('sanitizeHtml()', 'Strips HTML tags + encodes special chars'),
    ('hash()', 'SHA-256 hash via Web Crypto API'),
    ('Store class', 'localStorage wrapper with optional encryption'),
]
table = doc.add_table(rows=len(fields), cols=2)
table.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate(fields):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
doc.add_paragraph('Line 209: Hardcoded OpenRouter API key fallback — overridden by admin dashboard via rmh_bot_apikey localStorage.')

doc.add_heading('6.3 embed.js', level=2)
doc.add_paragraph('Chatbot widget loaded with defer. Key functions:')
efields = [
    ('formatBotMessage()', 'Rich formatting: **bold** -> <strong>, *italic* -> <em>, bullet lists -> <ul><li>, double newlines -> <p>'),
    ('embToggle()', 'Toggle chatbot open/closed'),
    ('embScroll()', 'Auto-scroll chat to bottom'),
    ('embToast()', 'Inline toast notification (replaces alert())'),
    ('Meeting tab', 'Date/time picker, form validation, localStorage save'),
    ('OpenRouter call', 'Configurable model, prompt injection detection, error handling'),
]
table = doc.add_table(rows=len(efields), cols=2)
table.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate(efields):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('6.4 chef-recipes.js', level=2)
doc.add_paragraph('28KB data file loaded with defer. Exposes C namespace:')
cfields = [
    ('C.CHEFS', '8 chef profiles (Sanjeev Kapoor, Ranveer Brar, Kunal Kapur, Tarla Dalal, Nisha Madhulika, VahChef, Hebbar\'s Kitchen, Swasthi\'s Recipes)'),
    ('C.CHEF_RECIPES', '9 full recipes with attribution and source URLs'),
    ('C.COMMODITY_PRODUCTS', '14 commodity items with INR default prices'),
    ('C.estimateRecipeCost()', 'Heuristic cost estimator based on ingredients'),
    ('C.fetchCommodityPrices()', 'Build-time API stub for data.gov.in), max 3 submissions per hour per browser'),
]
table = doc.add_table(rows=len(cfields), cols=2)
table.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate(cfields):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_page_break()

# ── 7. Configuration Guide ──
doc.add_heading('7. Configuration Guide', level=1)

doc.add_heading('7.1 API Key Setup', level=2)
steps = [
    'Open admin dashboard (/admin.html) in browser',
    'Login with default password: admin123',
    'Navigate to API Keys tab',
    'Paste your OpenRouter API key (get one at https://openrouter.ai/keys)',
    'Click Save',
    'Select a model from the dropdown (e.g., openrouter/free, google/gemini-2.5-pro-exp:free)',
    'Click "Test API Key" to verify',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('7.2 Changing Admin Password', level=2)
steps = [
    'Go to Password tab in admin dashboard',
    'Enter current password (default: admin123)',
    'Enter new password',
    'Confirm new password',
    'Click "Change Password"',
    'The next login will require the new password',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('7.3 API Key Storage Chain', level=2)
doc.add_paragraph(
    '1. Admin dashboard saves API key as encrypted copy (rmh_admin_apiKey) + plaintext copy (rmh_bot_apikey)\n'
    '2. utils.js (sync, loads before embed.js) reads rmh_bot_apikey and sets window.__RMH.API_KEY\n'
    '3. embed.js uses window.__RMH.API_KEY for OpenRouter API calls, with hardcoded fallback'
)

doc.add_heading('7.4 Theme Configuration', level=2)
doc.add_paragraph('In the Theming tab:')
items = [
    'Dark/Light mode toggle',
    'Accent color picker (default: #d4a017 — royal gold)',
    'Font selector (from Google Fonts library)',
    'Border radius slider (0-30px, default: 14px)',
    '5 preset themes: Royal (default), Midnight, Emerald, Ruby, Ocean',
]
for i in items:
    doc.add_paragraph(i, style='List Bullet')

doc.add_page_break()

# ── 8. Deployment Guide ──
doc.add_heading('8. Deployment Guide', level=1)

doc.add_heading('8.1 Netlify Deployment (Recommended)', level=2)
doc.add_paragraph('Method 1 — Git Import:')
steps = [
    'Push the project folder to a Git repository (GitHub/GitLab/Bitbucket)',
    'Go to https://app.netlify.com → Add new site → Import from Git',
    'Select the repository → Deploy (no build command needed, publish directory is .)',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_paragraph('Method 2 — Drag and Drop:')
steps = [
    'Go to https://app.netlify.com → Sites',
    'Drag and drop the entire client folder onto the browser window',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('8.2 Vercel Deployment', level=2)
steps = [
    'Push to a Git repository',
    'Go to https://vercel.com → Add New → Project',
    'Import the repository → Deploy (auto-detects static config from vercel.json)',
    'Alternatively: run "npx vercel --prod" from the project directory',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('8.3 Post-Deployment Checklist', level=2)
checks = [
    'Open https://your-site.netlify.app — verify Hero, cards, recipes render correctly',
    'Open https://your-site.netlify.app/admin.html — login works (default: admin123)',
    'Go to API Keys → Paste OpenRouter key → Save → Test',
    'Go to Password → Change password from default',
    'Open chatbot → Type "Schedule Meeting" → Fill form → Submit',
    'Go to Meetings tab in admin → Verify meeting appears',
    'Open a RecipeDB static page and verify it renders with JSON-LD',
    'Test on mobile — responsive layout, no false-positive anti-bot triggers',
    'Verify CSP headers via browser DevTools or curl',
]
for i, c in enumerate(checks, 1):
    doc.add_paragraph(f'{i}. {c}')

doc.add_page_break()

# ── 9. Admin Dashboard Guide ──
doc.add_heading('9. Admin Dashboard Guide', level=1)

doc.add_heading('9.1 Access', level=2)
doc.add_paragraph('Navigate to /admin.html. Default password: admin123.')

doc.add_heading('9.2 Tab Reference', level=2)

tab_details = [
    ('Overview', 'Displays 4 stat cards: Total Recipes (seed + chef + user), API Calls, Total Leads, Active Meetings. Bottom section shows recent activity log.'),
    ('API Keys', 'Encrypted storage for OpenRouter API key. Features: Show/Hide toggle, Test API Key validation, Model selector dropdown, Rotate key functionality. Saves plaintext copy to rmh_bot_apikey for the chatbot widget.'),
    ('Password', 'Change password form with 3 fields: current password, new password, confirm new password. Displays current password hash.'),
    ('Branding', 'Form to customize: business name, tagline (EN/TE), address (EN/TE), phone, WhatsApp, email. Mutton cuts: add/edit/delete product names. Custom CSS textarea for injecting styles.'),
    ('Theming', 'Dark/Light mode toggle, accent color picker, font selector (Google Fonts), border radius slider. 5 preset themes with instant preview.'),
    ('RecipeDB', 'Full CRUD interface for recipes. Features: search bar, filter by type (all/user/chef). Add new recipe form, edit existing, delete individual, bulk delete user recipes. Export/import JSON. User-submitted recipes (prefixed with user_) are editable; chef recipes view-only.'),
    ('Meetings', 'Table of B2B meeting requests with columns: Name, Phone, Email, Date, Time, Notes, Status, Actions. Filter by status (all/pending/confirmed/completed/cancelled). Search by name/phone/email. Actions: Confirm (pending->confirmed), Complete (confirmed->completed), Cancel, Reopen (cancelled->pending).'),
    ('Data Management', 'Encrypted backup: export all data encrypted with admin password. Import: restore from encrypted backup file. Local snapshots: create/view last 5 snapshots. CSV/JSON import/export for: Recipes, Meetings, Leads, Reviews & Ratings. Danger Zone: Clear User Recipes, Reset Analytics, Factory Reset (deletes ALL data).'),
    ('API Usage', '4 stat cards: Total Page Views, API Calls, Total Leads, Avg Views/Day. Leads table: Name, Source (chatbot/form with color badge), Date, Recipe. API Calls history: Date, Model (code styled), Query (truncated to 80 chars).'),
]
for name, desc in tab_details:
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ── 10. User Manual — Admin Operations ──
doc.add_heading('10. User Manual — Admin Operations', level=1)
doc.add_paragraph(
    'This section provides step-by-step instructions for all admin operations. '
    'The admin dashboard is accessible at /admin.html and requires password authentication.'
)

doc.add_heading('10.1 Logging In', level=2)
steps = [
    'Navigate to https://your-site.com/admin.html in your browser',
    'Enter your admin password (default: admin123)',
    'Click "Unlock Dashboard"',
    'On first login, change the default password immediately via the Password tab',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('10.2 Managing API Key', level=2)
steps = [
    'After logging in, click "API Keys" in the sidebar',
    'Paste your OpenRouter API key into the "API Key" field',
    '(Optional) Select a different AI model from the dropdown',
    'Click "Save API Key" — a success toast will confirm',
    'Click "Test API Key" to verify the key works',
    'Use the eye icon to toggle key visibility',
    'To rotate: click "Rotate", confirm, then enter a new key',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('10.3 Changing Password', level=2)
steps = [
    'Click "Password" in the sidebar',
    'Enter your current password',
    'Enter your new password (minimum 6 characters)',
    'Re-enter the new password to confirm',
    'Click "Change Password"',
    'A success toast confirms the change',
    'The current password hash is displayed for reference',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('10.4 Customizing Branding', level=2)
steps = [
    'Click "Branding" in the sidebar',
    'Edit any field: Business Name, Tagline (English), Tagline (Telugu), Address (English), Address (Telugu), Phone, WhatsApp, Email',
    'Each field auto-saves when you type (debounced)',
    'Mutton Cuts section: type a cut name and click "Add Cut" to add, click the X icon next to a cut to remove it',
    'Custom CSS: write CSS rules in the textarea (e.g., .hero { background: red; }) — applied site-wide',
    'Changes reflect immediately on the dashboard',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('10.5 Applying Themes', level=2)
steps = [
    'Click "Theming" in the sidebar',
    'Toggle Dark/Light mode using the switch',
    'Pick an accent color using the color picker',
    'Select a font from the dropdown (loaded from Google Fonts)',
    'Adjust border radius using the slider (0-30px)',
    'Or click one of the 5 preset theme buttons for instant styling',
    'All changes preview in real-time',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('10.6 Managing Recipes (RecipeDB CRUD)', level=2)
doc.add_heading('Viewing & Searching', level=3)
steps = [
    'Click "RecipeDB" in the sidebar',
    'All recipes are displayed in a table with columns: Name, Author, Type (seed/chef/user), Diet, Cuisine',
    'Use the search bar to filter by name, author, or cuisine',
    'Use the dropdown to filter by type: All, User Recipes, or Chef Recipes',
    'Click column headers to sort',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Adding a Recipe', level=3)
steps = [
    'Click "Add Recipe" button above the table',
    'Fill in the recipe form: name, author, category, cuisine, meal type, diet, budget, cost estimate, difficulty, prep time, cook time, servings, image (emoji), dietary restrictions',
    'Enter ingredients (one per line)',
    'Enter steps (one per line)',
    'Enter tips (one per line)',
    'Enter equipment (comma-separated)',
    'Click "Save" to add the recipe',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Editing a Recipe', level=3)
steps = [
    'Find the recipe in the table',
    'Click the edit (pencil) icon in the Actions column',
    'Modify any fields in the form that appears',
    'Click "Save Changes"',
    'Note: Chef recipes are view-only and cannot be edited',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Deleting a Recipe', level=3)
steps = [
    'Click the delete (trash) icon next to a recipe',
    'Confirm the deletion in the dialog',
    'Only user recipes can be deleted individually',
    'To delete all user recipes: use "Clear User Recipes" in Data Management > Danger Zone',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('10.7 Managing Meetings', level=2)
steps = [
    'Click "Meetings" in the sidebar',
    'All meeting requests appear in a table with: Name, Phone, Email, Date, Time, Notes, Status, Actions',
    'Use the search bar to find by name, phone, or email',
    'Use the status filter dropdown to view: All, Pending, Confirmed, Completed, Cancelled',
    'Actions per meeting:',
]
substeps = [
    'Green checkmark → Confirm (moves from pending to confirmed)',
    'Double checkmark → Complete (moves from confirmed to completed)',
    'Red X → Cancel (cancels the meeting)',
    'Undo arrow → Reopen (moves from cancelled back to pending)',
]
for s in substeps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Flow', level=3)
doc.add_paragraph('pending → confirmed → completed (happy path)')
doc.add_paragraph('pending/corfirmed → cancelled (rejection path)')
doc.add_paragraph('cancelled → pending (reopen)')

doc.add_heading('10.8 Data Management', level=2)

doc.add_heading('Encrypted Backup', level=3)
steps = [
    'Click "Data Management" in the sidebar',
    'Under "Export All Data", enter your admin password',
    'Click "Export Encrypted Backup" — downloads a .rmhbak file',
    'To restore: select the .rmhbak file, enter the same password used during export',
    'Click "Import & Restore" — all data is restored. Reload the page.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Local Snapshots', level=3)
steps = [
    'Click "Snapshot Now" to save a plain JSON summary (no API key)',
    'Last 5 snapshots appear below the button with timestamp and summary',
    'Maximum 20 snapshots are retained (oldest auto-removed)',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Export / Import CSV & JSON', level=3)
doc.add_paragraph('Located in the "Export / Import (CSV & JSON)" card below the main data management grid.')

doc.add_paragraph('Exporting:')
steps = [
    'Recipes: Click "JSON" or "CSV" under Recipes — downloads all user recipes',
    'Meetings: Click "JSON" or "CSV" under Meetings — downloads all meeting requests',
    'Leads: Click "JSON" or "CSV" under Leads — downloads all captured leads',
    'Reviews & Ratings: Click "JSON" under Reviews for text reviews, "Ratings" for appreciation data',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_paragraph('Importing (Recipes & Meetings):')
steps = [
    'Click the "JSON" or "CSV" button under Import for the data type you want to import',
    'Select the file from your computer',
    'The data is appended to existing records (no duplicates removed)',
    'A toast confirms how many records were imported',
    'CSV format: first row must be headers matching the field names. Arrays (ingredients, steps, etc.) use semicolons as separators.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Danger Zone', level=3)
steps = [
    'Clear User Recipes — deletes all user-submitted recipes (seed & chef recipes are preserved)',
    'Reset Analytics — clears all page views, API calls, and lead data',
    'Factory Reset All Data — deletes EVERYTHING: passwords, API keys, all recipes, analytics, backups, themes, branding. Page reloads automatically.',
]
doc.add_paragraph('⚠️ WARNING: Factory reset is irreversible. Two confirmation dialogs protect against accidental use.')
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('10.9 Monitoring API Usage', level=2)
steps = [
    'Click "API Usage" in the sidebar',
    'Four stat cards show: Total Page Views, API Calls, Total Leads, Avg Views/Day',
    'Leads table shows: Name, Source (chatbot/form with color-coded badge), Date, Recipe (if applicable)',
    'API Calls table shows: Date, Model name, Query preview (first 80 characters)',
    'Data refreshes when you visit the tab',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('10.10 Chatbot Operations (User-Facing)', level=2)
doc.add_paragraph('The chatbot is visible on the main site (bottom-right widget). Users can:')
items = [
    'Chat Tab — Ask questions about mutton supply, pricing, recipes, etc. AI-powered with canned fallback.',
    'Quote Tab — Request a B2B quote by filling name, phone, email, quantity, notes.',
    'Order Tab — Place an order by providing name, phone, email, cut selection, quantity, delivery date.',
    'Meeting Tab — Schedule a B2B consultation. Fill: name, phone, email, preferred date, time, notes. Saves to rmh_meetings with pending status.',
    'Policy Tab — View business policies: delivery, payment, quality assurance, cancellation.',
]
for i in items:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('10.11 Recipe Features (User-Facing)', level=2)
items = [
    'Browse recipes in the grid — filter by search, diet, budget, cuisine, meal type, category',
    'Click a recipe card to open the modal with full details',
    'Submit a recipe via the "Submit Your Recipe" form — rate-limited to 3/hour',
    'Rate recipes (1-5 stars) and mark "I Cooked This!"',
    'Post text reviews with your name and feedback',
    'Share recipes via WhatsApp, Facebook, X, Telegram, Email, Copy Link, or open the full RecipeDB page',
    'Filter by dietary restrictions: gluten-free, dairy-free, nut-free, egg-free, vegan',
]
for i in items:
    doc.add_paragraph(i, style='List Bullet')

doc.add_page_break()

# ── 12. localStorage Reference ──
doc.add_heading('12. localStorage Reference', level=1)
doc.add_paragraph('All data is stored client-side in the browser\'s localStorage. Key reference:')

table = doc.add_table(rows=18, cols=2)
table.style = 'Light Grid Accent 1'
locs = [
    ('rmh_admin_hash', 'SHA-256 hash of admin password'),
    ('rmh_admin_apiKey', 'OpenRouter API key (AES-256-GCM encrypted)'),
    ('rmh_bot_apikey', 'OpenRouter API key (plaintext, for chatbot use)'),
    ('rmh_admin_model', 'Selected AI model (e.g., openrouter/free)'),
    ('rmh_admin_branding', 'Branding overrides (name, tagline, address, etc.)'),
    ('rmh_admin_cuts', 'Custom mutton cuts array'),
    ('rmh_admin_theme', 'Theme settings (mode, accent, font, radius)'),
    ('rmh_admin_customCss', 'Custom CSS rules'),
    ('rmh_admin_backups', 'Local backup snapshots array (max 20)'),
    ('rmh_recipes', 'User-submitted recipes array'),
    ('rmh_ratelimit', 'Rate limiter: recipe_submit timestamps array'),
    ('rmh_appreciation', 'Ratings & "I Cooked This" counts per recipe ID'),
    ('rmh_reviews', 'User text reviews per recipe ID'),
    ('rmh_meetings', 'B2B meeting requests with status tracking'),
    ('rmh_analytics', 'Analytics: pageViews[], apiCalls[], leads[]'),
    ('rmh_bot_apikey', 'Plaintext API key for bot (set by admin dashboard)'),
    ('rmh_admin_model', 'AI model selection'),
]
for i, (k, v) in enumerate(locs):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_page_break()

# ── 13. Build Scripts ──
doc.add_heading('13. Build Scripts', level=1)

doc.add_heading('13.1 generate-recipes.js', level=2)
doc.add_paragraph(
    'Parses both utils.js (BUSINESS.recipes) and chef-recipes.js (C.CHEF_RECIPES) arrays, '
    'then generates 22 static HTML files in the RecipeDB/ folder with per-recipe JSON-LD, '
    'Open Graph tags, chef badges, share buttons, and prev/next navigation.'
)
doc.add_paragraph('Run: node generate-recipes.js')
doc.add_paragraph('Must be re-run after any recipe changes (add/edit/delete).')

doc.add_heading('13.2 fetch-commodity-prices.js', level=2)
doc.add_paragraph(
    'Node.js build script that fetches live mandi prices from the data.gov.in API. '
    'Includes 14 commodity defaults in INR as fallback.'
)
doc.add_paragraph('Run: node fetch-commodity-prices.js --output prices.json')

doc.add_heading('13.3 generate_doc.py', level=2)
doc.add_paragraph('This script. Generates this Word documentation file.')
doc.add_paragraph('Run: python generate_doc.py')

doc.add_page_break()

# ── 14. Screenshot Checklist ──
doc.add_heading('14. Screenshot Checklist', level=1)
doc.add_paragraph('Capture the following screenshots for documentation or client review:')

doc.add_heading('14.1 Main Site (index.html)', level=2)
ss_main = [
    'Hero — Brand badges, CTA buttons, district badges at top',
    'How We Work — 5-step process cards',
    'B2B Supply — 6 service cards',
    'Why Us — Indigenous highlight + 9 benefit cards',
    'Testimonials — 6 client reviews',
    'Recipe Grid — Recipe cards with filters and search bar',
    'Recipe Modal — Full recipe view with ingredients, instructions, equipment, rating, reviews, share buttons',
    'Contact — Bilingual address, OpenStreetMap, enquiry form',
    'Chatbot — AI chatbot open with message interaction',
    'Meeting Form — Meeting scheduling tab in chatbot',
]
for s in ss_main:
    p = doc.add_paragraph(s, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)

doc.add_heading('14.2 RecipeDB (recipedb.html)', level=2)
ss_recipedb = [
    'Recipe List — Full recipe grid with filters',
    'Recipe Detail — Static recipe page with JSON-LD structured data',
]
for s in ss_recipedb:
    p = doc.add_paragraph(s, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)

doc.add_heading('14.3 Admin Dashboard (admin.html)', level=2)
ss_admin = [
    'Login — Password login screen',
    'Overview — Stats cards and recent activity',
    'API Keys — Encrypted API key management',
    'Password — Change password form',
    'Branding — Business info, cuts, custom CSS',
    'Theming — Dark/light toggle, accent color, font, radius, presets',
    'RecipeDB — Recipe CRUD table with search',
    'Meetings — Meeting requests table with status actions',
    'Data Management — Backup/restore, export/import, CSV/JSON tools, danger zone',
    'API Usage — Analytics stats, leads table, API calls history',
]
for s in ss_admin:
    p = doc.add_paragraph(s, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)

# ── Save ──
output_path = 'C:\\Users\\saikarun\\Downloads\\client\\Royal_Mutton_Hub_Documentation.docx'
doc.save(output_path)
print(f'Documentation saved to: {output_path}')
